#!/usr/bin/env python3
"""Extract text and basic metadata from one uploaded resume without semantic invention."""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

SUPPORTED = {".pdf", ".docx", ".txt", ".md", ".markdown"}
MIN_NON_WS = 80


class ResumeError(Exception):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalize(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    output, previous_blank = [], False
    for line in lines:
        if not line:
            if not previous_blank:
                output.append("")
            previous_blank = True
        else:
            output.append(line)
            previous_blank = False
    return "\n".join(output).strip()


def extract_pdf(path: Path) -> tuple[str, str, int, list[str]]:
    warnings: list[str] = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            try:
                if reader.decrypt("") == 0:
                    raise ResumeError("PDF_UNREADABLE", "PDF is password-protected.")
            except Exception as exc:
                raise ResumeError("PDF_UNREADABLE", "PDF is password-protected.") from exc
        page_count = len(reader.pages)
    except ResumeError:
        raise
    except Exception as exc:
        raise ResumeError("PDF_UNREADABLE", f"Unable to open PDF: {exc}") from exc

    if shutil.which("pdftotext"):
        proc = subprocess.run(
            ["pdftotext", "-layout", "-enc", "UTF-8", str(path), "-"],
            capture_output=True,
            text=True,
        )
        if proc.returncode == 0:
            return proc.stdout, "pdftotext-layout", page_count, warnings
        warnings.append(f"pdftotext failed: {proc.stderr.strip()}")

    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
        return text, "pypdf", page_count, warnings
    except Exception as exc:
        raise ResumeError("PDF_UNREADABLE", f"Unable to extract PDF text: {exc}") from exc


def extract_docx(path: Path) -> tuple[str, str, int, list[str]]:
    try:
        from docx import Document
        document = Document(str(path))
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(blocks), "python-docx", 0, []
    except Exception as exc:
        raise ResumeError("DOCX_UNREADABLE", f"Unable to read DOCX: {exc}") from exc


def extract_plain(path: Path) -> tuple[str, str, int, list[str]]:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding), f"text-{encoding}", 0, []
        except UnicodeDecodeError:
            continue
    raise ResumeError("TEXT_UNREADABLE", "Unable to decode text file as UTF-8 or GB18030.")


def basic_contacts(text: str) -> dict[str, list[str]]:
    emails = sorted(set(re.findall(r"(?i)\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", text)))
    phones = sorted(set(re.findall(r"(?<!\d)(?:\+?86[- ]?)?1[3-9]\d{9}(?!\d)", text)))
    urls = sorted(set(re.findall(r"https?://[^\s<>()\[\]{}]+", text)))
    return {"emails": emails, "phones": phones, "urls": urls}


def section_hints(text: str) -> dict[str, list[str]]:
    patterns = {
        "education": r"教育背景|教育经历|EDUCATION",
        "skills": r"专业技能|技能|SKILLS|TECHNICAL",
        "projects": r"项目经历|科研经历|PROJECT|RESEARCH",
        "experience": r"实习经历|工作经历|EXPERIENCE|EMPLOYMENT",
        "publications": r"论文|科研成果|PUBLICATION",
        "awards": r"获奖|荣誉|AWARD|HONOR",
    }
    lines = text.splitlines()
    result: dict[str, list[str]] = {}
    for key, pattern in patterns.items():
        hits = []
        for index, line in enumerate(lines):
            if re.search(pattern, line, flags=re.I):
                hits.append("\n".join(lines[index : min(len(lines), index + 8)]))
        if hits:
            result[key] = hits[:5]
    return result


def atomic_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile("w", encoding="utf-8", dir=path.parent, delete=False) as stream:
        json.dump(payload, stream, ensure_ascii=False, indent=2)
        temporary = Path(stream.name)
    temporary.replace(path)


def error_payload(path: Path | None, code: str, message: str) -> dict:
    return {
        "schema_version": 1,
        "status": "error",
        "source": {"file_name": path.name if path else "", "extension": path.suffix.lower() if path else ""},
        "stats": {"character_count": 0, "non_whitespace_count": 0, "line_count": 0},
        "contacts": {"emails": [], "phones": [], "urls": []},
        "section_hints": {},
        "warnings": [],
        "raw_text": "",
        "error": {"code": code, "message": message},
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("resume")
    parser.add_argument("--output", required=True)
    parser.add_argument("--text-output")
    parser.add_argument("--min-non-whitespace", type=int, default=MIN_NON_WS)
    args = parser.parse_args()

    resume = Path(args.resume).expanduser().resolve()
    output = Path(args.output).expanduser().resolve()
    text_output = Path(args.text_output).expanduser().resolve() if args.text_output else None

    try:
        if not resume.exists() or not resume.is_file():
            raise ResumeError("RESUME_REQUIRED", f"Resume file does not exist: {resume}")
        if resume.stat().st_size == 0:
            raise ResumeError("EMPTY_FILE", "Resume file is empty.")
        extension = resume.suffix.lower()
        if extension not in SUPPORTED:
            raise ResumeError("UNSUPPORTED_FORMAT", f"Unsupported resume format: {extension or '[none]'}. Use PDF, DOCX, TXT, or Markdown.")

        if extension == ".pdf":
            raw, parser_name, pages, warnings = extract_pdf(resume)
        elif extension == ".docx":
            raw, parser_name, pages, warnings = extract_docx(resume)
        else:
            raw, parser_name, pages, warnings = extract_plain(resume)
        text = normalize(raw)
        non_whitespace = len(re.sub(r"\s", "", text))
        status = "success" if non_whitespace >= args.min_non_whitespace else "needs_ocr"
        error = None
        if status == "needs_ocr":
            warnings.append(f"Only {non_whitespace} non-whitespace characters were extracted; OCR or a text-searchable resume is required.")
            error = {"code": "OCR_REQUIRED", "message": warnings[-1]}

        payload = {
            "schema_version": 1,
            "status": status,
            "source": {
                "file_name": resume.name,
                "extension": extension,
                "sha256": sha256(resume),
                "size_bytes": resume.stat().st_size,
                "parsed_at_utc": dt.datetime.now(dt.timezone.utc).isoformat(),
                "parser": parser_name,
                "page_count": pages,
            },
            "stats": {
                "character_count": len(text),
                "non_whitespace_count": non_whitespace,
                "line_count": len(text.splitlines()),
            },
            "contacts": basic_contacts(text),
            "section_hints": section_hints(text),
            "warnings": warnings,
            "raw_text": text,
            "error": error,
        }
        atomic_json(output, payload)
        if text_output:
            text_output.parent.mkdir(parents=True, exist_ok=True)
            text_output.write_text(text, encoding="utf-8")
        print(json.dumps({"status": status, "output": str(output), "sha256": payload["source"]["sha256"], "non_whitespace_count": non_whitespace}, ensure_ascii=False))
        return 0 if status == "success" else 2
    except ResumeError as exc:
        atomic_json(output, error_payload(resume, exc.code, exc.message))
        print(json.dumps({"status": "error", "code": exc.code, "message": exc.message, "output": str(output)}, ensure_ascii=False), file=sys.stderr)
        return 1
    except Exception as exc:
        atomic_json(output, error_payload(resume, "UNEXPECTED_PARSE_ERROR", str(exc)))
        print(json.dumps({"status": "error", "code": "UNEXPECTED_PARSE_ERROR", "message": str(exc), "output": str(output)}, ensure_ascii=False), file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
